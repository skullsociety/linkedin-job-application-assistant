import sqlite3
import tempfile
import unittest
from contextlib import closing
from pathlib import Path

from openpyxl import load_workbook
from docx import Document

from job_assistant.__main__ import match_resume
from job_assistant.config import Settings
from job_assistant.database import JobRepository
from job_assistant.exporter import export_jobs
from job_assistant.models import Job


class DatabaseAndExportTests(unittest.TestCase):
    def test_source_defaults_to_linkedin_and_round_trips(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with JobRepository(Path(directory) / "jobs.sqlite3") as repository:
                saved = repository.upsert(Job("Engineer", "Example", "https://example.test/jobs/source"))
                raw_source = repository.connection.execute(
                    "SELECT source FROM jobs WHERE id=?", (saved.id,)
                ).fetchone()[0]
                source_column = next(
                    row for row in repository.connection.execute("PRAGMA table_info(jobs)") if row[1] == "source"
                )
                self.assertEqual(saved.source, "linkedin")
                self.assertEqual(raw_source, "linkedin")
                self.assertEqual(source_column[2], "TEXT")
                self.assertEqual(source_column[3], 1)
                self.assertEqual(source_column[4], "'linkedin'")

                updated = repository.upsert(Job(
                    "Engineer", "Example", "https://example.test/jobs/source", source="indeed"
                ))
                self.assertEqual(updated.source, "indeed")

    def test_existing_database_adds_default_source_without_losing_jobs(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            database_path = Path(directory) / "jobs.sqlite3"
            with closing(sqlite3.connect(database_path)) as connection:
                connection.execute(
                    """CREATE TABLE jobs (
                        id INTEGER PRIMARY KEY,
                        title TEXT NOT NULL,
                        company TEXT NOT NULL,
                        url TEXT NOT NULL UNIQUE,
                        date_found TEXT NOT NULL,
                        match_score INTEGER,
                        notes TEXT,
                        status TEXT NOT NULL DEFAULT 'saved',
                        submission_approved_at TEXT,
                        created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
                    )"""
                )
                connection.execute(
                    "INSERT INTO jobs (title, company, url, date_found) VALUES (?, ?, ?, ?)",
                    ("Legacy Engineer", "Example", "https://example.test/jobs/legacy", "2026-08-24"),
                )
                connection.commit()

            with JobRepository(database_path) as repository:
                saved = repository.list()[0]
                self.assertEqual(saved.title, "Legacy Engineer")
                self.assertEqual(saved.source, "linkedin")
                self.assertEqual(repository.connection.execute("SELECT COUNT(*) FROM jobs").fetchone()[0], 1)

    def test_save_and_export_job(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            tmp_path = Path(directory)
            repository = JobRepository(tmp_path / "jobs.sqlite3")
            saved = repository.upsert(Job("Engineer", "Example Co", "https://example.test/jobs/1", location="Singapore", salary="SGD 100000", platform="Example", match_score=85))
            repository.save_match(
                saved.id or 0,
                score=85,
                matching_skills="python, excel",
                missing_skills="sql, tableau",
                reason="Test match.",
                recommendation="apply",
            )
            repository.record_submission_approval(saved.id or 0)
            output = export_jobs(repository.list(), tmp_path / "tracker.xlsx")
            sheet = load_workbook(output).active
            self.assertEqual(sheet.max_row, 4)
            self.assertEqual(sheet.cell(4, 2).value, "Example")
            self.assertEqual(sheet.cell(4, 4).value, "Engineer")
            self.assertEqual(sheet.cell(4, 9).value, "sql, tableau")
            self.assertEqual(sheet.cell(4, 10).value, "ready_for_manual_submit")
            self.assertIsNone(sheet.auto_filter.ref)
            self.assertEqual(sheet.tables["JobApplications"].ref, "A3:L4")
            repository.connection.close()

    def test_export_removes_duplicate_urls(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = export_jobs([
                Job("Engineer", "Example", "https://example.test/jobs/duplicate"),
                Job("Engineer II", "Example", "https://example.test/jobs/duplicate"),
            ], Path(directory) / "tracker.xlsx")
            self.assertEqual(load_workbook(output).active.max_row, 4)

    def test_export_overwrites_the_same_tracker_file(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "tracker.xlsx"
            export_jobs([Job("First role", "Example", "https://example.test/jobs/first")], output)
            export_jobs([Job("Updated role", "Example", "https://example.test/jobs/updated")], output)
            sheet = load_workbook(output).active
            self.assertEqual(sheet.max_row, 4)
            self.assertEqual(sheet.cell(4, 4).value, "Updated role")

    def test_resume_matching_refreshes_the_tracker_automatically(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            resume = root / "resume.docx"
            document = Document()
            document.add_paragraph("Python and SQL")
            document.save(resume)
            settings = Settings(
                database_path=root / "jobs.sqlite3",
                export_path=root / "tracker.xlsx",
                log_path=root / "assistant.log",
                browser_user_data_dir=root / "browser",
                browser_channel="chrome",
                profile_path=root / "profile.json",
                tailored_resume_dir=root / "tailored",
                tailored_resume_threshold=100,
                resume_dir=root / "resumes",
                resume_path=None,
            )
            with JobRepository(settings.database_path) as repository:
                repository.upsert(Job("Engineer", "Example", "https://example.test/jobs/match", job_description="Python and SQL"))
                match_resume(repository, resume, None, settings, reporter=lambda _: None)
            sheet = load_workbook(settings.export_path).active
            self.assertEqual(sheet.cell(4, 4).value, "Engineer")

    def test_update_rejects_invalid_follow_up_date(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with JobRepository(Path(directory) / "jobs.sqlite3") as repository:
                job = repository.upsert(Job("Engineer", "Example", "https://example.test/jobs/date"))
                with self.assertRaisesRegex(ValueError, "YYYY-MM-DD"):
                    repository.update(job.id or 0, score=None, notes=None, status=None, follow_up_date="next week")

    def test_linkedin_tracking_urls_save_as_one_job(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with JobRepository(Path(directory) / "jobs.sqlite3") as repository:
                first = repository.upsert(Job("Data Scientist", "GoTo", "https://www.linkedin.com/jobs/search-results/?currentJobId=4378574897&eBP=first"))
                second = repository.upsert(Job("Data Scientist", "GoTo", "https://www.linkedin.com/jobs/search-results/?currentJobId=4378574897&eBP=second"))
                self.assertEqual(first.id, second.id)
                self.assertEqual(len(repository.list()), 1)
                self.assertEqual(first.url, "https://www.linkedin.com/jobs/view/4378574897/")

    def test_capture_metadata_and_applied_events_are_preserved(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with JobRepository(Path(directory) / "jobs.sqlite3") as repository:
                first = repository.upsert(Job(
                    "Cloud Engineer",
                    "Example",
                    "https://www.linkedin.com/jobs/view/12345/",
                    linkedin_job_id="12345",
                    company_url="https://www.linkedin.com/company/example/",
                    application_method="Easy Apply",
                    employment_type="Full-time",
                    description_hash="abc123",
                    first_seen_at="2026-08-21T10:00:00+08:00",
                    last_seen_at="2026-08-21T10:00:00+08:00",
                ))
                repeated = repository.upsert(Job(
                    "Cloud Engineer",
                    "Example",
                    "https://www.linkedin.com/jobs/view/12345/",
                    linkedin_job_id="12345",
                    application_method="Easy Apply",
                    first_seen_at="2026-08-21T10:05:00+08:00",
                    last_seen_at="2026-08-21T10:05:00+08:00",
                ))
                self.assertEqual(repeated.id, first.id)
                self.assertEqual(repeated.seen_count, 2)
                applied = repository.set_applied(first.id or 0, True)
                self.assertTrue(applied.applied)
                self.assertEqual(applied.status, "submitted_manually")
                self.assertIsNotNone(applied.applied_at)
                self.assertEqual(
                    [event.event_type for event in repository.events(first.id or 0)],
                    ["CAPTURED", "APPLICATION_SUBMITTED"],
                )
                unmarked = repository.set_applied(first.id or 0, False)
                self.assertFalse(unmarked.applied)
                self.assertIsNone(unmarked.applied_at)
                self.assertEqual(unmarked.status, "saved")
                self.assertEqual(repository.events(first.id or 0)[-1].event_type, "APPLICATION_UNMARKED")

    def test_delete_all_removes_saved_jobs(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with JobRepository(Path(directory) / "jobs.sqlite3") as repository:
                repository.upsert(Job("Engineer", "Example", "https://example.test/jobs/clear"))
                repository.upsert(Job("Analyst", "Example", "https://example.test/jobs/clear-2"))
                self.assertEqual(repository.delete_all(), 2)
                self.assertEqual(repository.list(), [])
