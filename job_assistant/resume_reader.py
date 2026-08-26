from __future__ import annotations

from pathlib import Path


SUPPORTED_RESUME_SUFFIXES = frozenset({".pdf", ".docx"})


def latest_resume(directory: Path) -> Path:
    """Return the most recently modified PDF or DOCX directly inside a resume folder."""
    if not directory.is_dir():
        raise FileNotFoundError(f"Resume folder not found: {directory}")
    candidates = [path for path in directory.iterdir() if path.is_file() and path.suffix.casefold() in SUPPORTED_RESUME_SUFFIXES]
    if not candidates:
        raise FileNotFoundError(f"No PDF or DOCX resume found in: {directory}")
    return max(candidates, key=lambda path: (path.stat().st_mtime_ns, path.name.casefold()))


def read_resume(path: Path) -> str:
    """Extract resume text locally. The resume is never uploaded or sent to a service."""
    if not path.is_file():
        raise FileNotFoundError(f"Resume not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError("Resume must be a PDF or DOCX file.")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ValueError("The PDF is password-protected. Save an unprotected copy and try again.")
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return _require_text(text, path)


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return _require_text("\n".join(paragraphs + table_cells), path)


def _require_text(text: str, path: Path) -> str:
    cleaned = text.strip()
    if not cleaned:
        raise ValueError(f"No readable text was found in {path.name}. Use a text-based PDF or DOCX.")
    return cleaned
